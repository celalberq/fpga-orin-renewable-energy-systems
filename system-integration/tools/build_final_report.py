from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "system-integration" / "demo_evidence" / "final_report"
TMP_DIR = OUT_DIR / "_build"
OUT_FILE = OUT_DIR / "fpga_orin_renewable_energy_systems_final_report_2026_08_11.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "18212A"
MUTED = "65717C"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
GREEN = "167447"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_picture_alt(inline_shape, description):
    inline_shape._inline.docPr.set("descr", description)
    inline_shape._inline.docPr.set("title", description)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, color=INK, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run(rest)
    else:
        run = p.add_run(text)
        set_run(run)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run(run, size=9, color=MUTED, italic=True)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_metric_table(doc, metrics):
    table = doc.add_table(rows=2, cols=len(metrics))
    table.style = "Table Grid"
    widths = [9360 // len(metrics)] * len(metrics)
    widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for col, (label, value) in enumerate(metrics):
        set_cell_shading(table.cell(0, col), PALE_BLUE)
        p = table.cell(0, col).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(label.upper()), size=8.5, color=DARK_BLUE, bold=True)
        p = table.cell(1, col).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(value), size=14, color=INK, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_result_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [3000, 4860, 1500])
    headers = ("Test", "Measured result", "Status")
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), size=9.5, color=DARK_BLUE, bold=True)
    set_repeat_table_header(table.rows[0])
    for name, result, status in rows:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margins(cell, top=100, bottom=100)
        for index, text in enumerate((name, result, status)):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(
                p.add_run(text),
                size=9.3,
                color=GREEN if index == 2 else INK,
                bold=index in (0, 2),
            )
    set_table_geometry(table, [3000, 4860, 1500])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc, label, text, fill=PALE_BLUE, accent=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    set_cell_shading(table.cell(0, 0), fill)
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(label + "  "), size=10, color=accent, bold=True)
    set_run(p.add_run(text), size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def crop_charge_image():
    source = ROOT / "system-integration" / "demo_evidence" / "image" / "bess_real_charge_dashboard_2026-08-11.png"
    target = TMP_DIR / "bess_charge_dashboard_crop.png"
    with Image.open(source) as image:
        crop = image.crop((0, 0, image.width, min(820, image.height)))
        crop.save(target)
    return target


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, color, before, after in (
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    set_run(header.add_run("THREE-PROJECT RENEWABLE-ENERGY PLATFORM"), size=8.5, color=MUTED, bold=True)
    set_run(header.add_run("\tFINAL REPORT | 11 AUG 2026"), size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    # Page 1: memo masthead and executive summary.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("FINAL TECHNICAL REPORT"), size=10, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run("FPGA and Jetson Orin Renewable-Energy Systems"), size=23, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    set_run(p.add_run("Renewable sensing, two-axis control, signed storage telemetry, digital communication profiling, and live network monitoring"), size=12.5, color=MUTED)

    for label, value in (
        ("Platform", "Digilent Nexys Video + Jetson Orin Nano Super"),
        ("Integrated projects", "Solar Tracker and BESS | FPGA Digital Communications | Network Telemetry Dashboard"),
        ("Evidence date", "11 August 2026"),
        ("Validation state", "Live integrated pass"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(label + ": "), size=10.5, color=INK, bold=True)
        set_run(p.add_run(value), size=10.5, color=INK)

    add_callout(doc, "HEADLINE RESULT", "One continuous 368-packet run demonstrated physical two-axis tracking and real battery idle, charge, discharge, and final idle recovery through a 161-byte CRC/QPSK/OFDM profile with zero sequence gaps.")
    add_heading(doc, "Executive Summary", 1)
    add_body(doc, "This portfolio contains three independently scoped projects that also form a complete low-voltage renewable edge demonstrator. Four light sensors are sampled through an MCP3208 and Nexys Video FPGA; a Jetson Orin estimates light direction and controls a physical two-axis solar-panel mount; and an INA226 measures a protected 1S Li-ion battery during supervised charging and fan-load discharge. The same live packet is profiled by a CRC/QPSK/OFDM communication project and monitored by a separate UDP dashboard project with packet validation, gap detection, protection state, and CSV evidence.")
    add_metric_table(doc, (("Tracking gain", "+24.141%"), ("Final frames", "368 / 368"), ("Sequence gaps", "0"), ("BESS modes", "3")))
    add_body(doc, "The three projects retain separate requirements, implementations, tests, and evidence while sharing a telemetry contract for integration: renewable sensing and control from Solar Tracker and BESS, deterministic communication framing from FPGA Digital Communications, and network observability from Network Telemetry Dashboard. The final run exercised all three while the measured battery state changed continuously from idle to charging, through a safe idle transition, to discharging, and back to idle.")

    doc.add_page_break()

    # Page 2: architecture and ownership.
    add_heading(doc, "1. System Architecture", 1)
    add_body(doc, "The architecture separates deterministic sampling and FPGA communication proof from Linux-side tracking, telemetry integration, and dashboard services.")
    diagram = doc.add_paragraph()
    diagram.alignment = WD_ALIGN_PARAGRAPH.CENTER
    diagram.paragraph_format.space_before = Pt(8)
    diagram.paragraph_format.space_after = Pt(10)
    diagram_text = (
        "4x KY-018 + panel divider\n"
        "        |\n"
        "MCP3208 -> Nexys Video UART -> Jetson Orin -> PCA9685 -> pan/tilt panel\n"
        "                                  |\n"
        "Battery + TP4056 -> INA226 --------+\n"
        "                                  |\n"
        "                         unified UDP telemetry\n"
        "                                  |\n"
        "                         CRC/QPSK/OFDM bridge\n"
        "                                  |\n"
        "                         dashboard + CSV evidence"
    )
    set_run(diagram.add_run(diagram_text), size=9.5, color=DARK_BLUE, font="Consolas")

    add_heading(doc, "Project Ownership", 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1800, 3780, 3780])
    for i, text in enumerate(("Project", "Primary responsibility", "Validated evidence")):
        set_cell_shading(table.cell(0, i), LIGHT)
        set_run(table.cell(0, i).paragraphs[0].add_run(text), size=9.5, color=DARK_BLUE, bold=True)
    set_repeat_table_header(table.rows[0])
    for project, responsibility, evidence in (
        ("Solar Tracker and BESS", "LDR and panel sensing; Orin tracking; PCA9685 servo actuation; INA226 BESS measurement; protection semantics", "Two-axis tracking, shading/recovery, charge/discharge, unified live run"),
        ("FPGA Digital Communications", "CRC frame construction; QPSK sizing; OFDM resource, pilot, padding, and cyclic-prefix accounting", "Nexys Video board tests and live 161-byte mixed-frame bridge"),
        ("Network Telemetry Dashboard", "UDP transport; packet validation; gap monitoring; dashboard state; CSV logging", "Healthy live link, zero-gap runs, reliability warning/recovery"),
    ):
        cells = table.add_row().cells
        for i, text in enumerate((project, responsibility, evidence)):
            set_run(cells[i].paragraphs[0].add_run(text), size=9.2, color=INK, bold=i == 0)
    set_table_geometry(table, [1800, 3780, 3780])
    add_heading(doc, "Unified Packet Profile", 2)
    add_metric_table(doc, (("Payload", "157 B"), ("CRC frame", "161 B"), ("QPSK", "644"), ("OFDM", "14")))
    add_body(doc, "The older tracker-only profile remains valid for comparison at 129 payload bytes, 133 frame bytes, 532 QPSK symbols, and 12 OFDM symbols. The larger unified profile adds real BESS fields without changing the zero-gap network behavior.")

    doc.add_page_break()

    # Page 3: implementation and physical evidence.
    add_heading(doc, "2. Implementation", 1)
    add_heading(doc, "Tracking and Actuation", 2)
    add_body(doc, "A fixed four-LDR cross provides an absolute two-axis target. The Orin applies sensor calibration, deadbands, sequential axis selection, bounded movement steps, and continuous reacquisition. Pan uses PCA9685 channel 15 and tilt uses channel 12. The replacement mechanism completed full two-axis tracking without the binding that damaged the earlier tilt assembly.")
    add_heading(doc, "Panel and Storage Measurement", 2)
    add_body(doc, "Panel voltage is sampled on MCP3208 CH0 through a safe divider and converted to load-power estimate using the known 1k-ohm resistor. The INA226 was moved to the battery path, calibrated with a 0.9501 voltage scale, and preserves signed current end to end: positive for charging and negative for discharging.")
    image = ROOT / "system-integration" / "reports" / "assets" / "demo_final_full.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    tracker_picture = p.add_run().add_picture(str(image), width=Inches(6.35))
    set_picture_alt(
        tracker_picture,
        "Integrated tracker evidence showing the live dashboard, Orin output, communication bridge, and physical solar panel.",
    )
    add_caption(doc, "Figure 1. Live tracker dashboard, Orin control output, communication bridge, and physical panel view.")
    add_callout(doc, "CONTROL RESULT", "The tracker reached locked state with physical servo output enabled; the controlled A/B experiment measured 24.141% higher mean estimated load power than the fixed-panel run.", fill="EEF7F2", accent=GREEN)

    doc.add_page_break()

    # Page 4: quantitative results and charge proof.
    add_heading(doc, "3. Measured Results", 1)
    add_result_table(doc, (
        ("Real LDR tracking", "368 valid packets; both axes moved; best error 0.012 degrees", "PASS"),
        ("Tracking A/B", "+24.141% mean-power gain; +24.229% energy gain", "PASS"),
        ("Shading recovery", "120 valid packets; warning at seq 61; recovered at seq 78", "PASS"),
        ("BESS discharge", "About 2.954 V, -231 mA, and -0.683 W with fan load", "PASS"),
        ("BESS charge", "3.835-3.842 V, +233 to +236 mA, +0.895 to +0.908 W", "PASS"),
        ("Unified discharge demo", "60/60 valid 161-byte frames; 0 gaps; real discharging state", "PASS"),
        ("Bidirectional BESS demo", "368/368 valid; 0 gaps; idle/charge/idle/discharge/idle", "PASS"),
    ))

    # Continue on the results page; the charging figure flows to the next page.
    add_heading(doc, "4. Continuous BESS Evidence", 1)
    add_body(doc, "The strongest final recording keeps tracker control, BESS measurement, communication profiling, UDP transport, dashboard state, and CSV logging active for the entire electrical transition. TP4056 USB-C and the fan load were never connected simultaneously.")
    add_metric_table(doc, (("Idle rows", "191"), ("Charging rows", "44"), ("Discharge rows", "133"), ("Locked rows", "213")))
    add_callout(doc, "STATE SEQUENCE", "seq 000-156 idle | 157-198 charging | 199-202 safe transition | 203-335 discharging | 336-367 final idle", fill="EEF7F2", accent=GREEN)
    add_heading(doc, "Real Charging Evidence", 2)
    charge_image = crop_charge_image()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    charge_picture = p.add_run().add_picture(str(charge_image), width=Inches(5.75))
    set_picture_alt(
        charge_picture,
        "INA226 charging dashboard showing approximately 3.84 volts, 0.24 amperes, 0.91 watts, and zero packet gaps.",
    )
    add_caption(doc, "Figure 2. INA226 battery charging dashboard: +0.24 A, 0.91 W, 40 valid packets, no invalid packets, and zero gaps.")
    add_body(doc, "The red TP4056 charging LED was illuminated during the supervised standalone charge test. In the continuous final demo, charging settled near +230 mA and approximately +0.83 W average; discharge reached -215.5 mA and -0.638 W. The final idle phase returned to approximately 3.50 V and near-zero current.")
    add_body(doc, "Measurement note: when both USB-C and the fan path were initially open, the INA226 bus-voltage register observed the open monitored path rather than a valid cell terminal voltage. Idle classification remained valid because signed current was near zero; no initial open-path voltage is used as a battery-voltage claim.")

    doc.add_page_break()

    # Page 6: safety, limitations, evidence, and conclusion.
    add_heading(doc, "5. Verification Boundaries", 1)
    add_heading(doc, "Safety Boundary", 2)
    add_body(doc, "The Orin and Nexys remain on their normal external supplies; the 1S battery does not power either board. The TP4056 USB input is disconnected while the fan load is connected because simultaneous charge/load power sharing has not been validated. Temporary mechanical battery connections are used only for supervised bench testing.")
    add_heading(doc, "Honest Limitations", 2)
    add_body(doc, "The tracker mechanics, LDR cross, panel mount, and battery terminals are prototype construction. Indoor panel power is low because the available source is a handheld LED, so panel power is an estimate from voltage across a known resistor rather than a claim of useful solar generation. The communication layer is CRC/QPSK/OFDM packet profiling and FPGA pipeline validation, not a complete over-the-air RF modem.")
    add_callout(doc, "OPTIONAL NEXT WORK", "Outdoor sunlight measurements, mechanically secure battery terminals and enclosure, a full FPGA IFFT/FFT modem or SDR link, and remote database/cloud deployment would extend the prototype but are not required for the validated portfolio claim.", fill="FFF8E8", accent=GOLD)
    add_heading(doc, "Primary Evidence", 2)
    for label, value in (
        ("Final bidirectional video", "system-integration/demo_evidence/video/unified_tracker_bess_bidirectional_final_demo_2026-08-11.mp4"),
        ("Final bridge CSV", "fpga-digital-communications/data/unified_tracker_bess_bidirectional_demo_bridge.csv"),
        ("Final dashboard CSV", "network-telemetry-dashboard/data/unified_tracker_bess_bidirectional_demo_dashboard.csv"),
        ("Final board-test record", "solar-tracker-bess/docs/board_test_25_unified_tracker_bess_bidirectional_demo.txt"),
        ("Charging board-test record", "solar-tracker-bess/docs/board_test_23_real_bess_charge_udp_dashboard.txt"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        set_run(p.add_run(label + ": "), size=9.5, color=INK, bold=True)
        set_run(p.add_run(value), size=9.2, color=MUTED, font="Consolas")
    add_heading(doc, "Conclusion", 2)
    add_body(doc, "The integrated demonstration connects physical sensing, real two-axis control, panel-power estimation, measured battery charging and discharge, FPGA communication validation, reliable UDP transport, and dashboard observability. Together, the three distinct projects provide portfolio evidence across FPGA design, embedded Linux, instrumentation, renewable energy, digital communications, networking, and data-driven verification.")

    doc.core_properties.title = "FPGA and Jetson Orin Renewable-Energy Systems"
    doc.core_properties.subject = "Final report for three integrated engineering projects"
    doc.core_properties.author = "Project Portfolio"
    doc.core_properties.keywords = "FPGA, Jetson Orin, solar tracker, BESS, INA226, QPSK, OFDM, UDP"
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build_document()
